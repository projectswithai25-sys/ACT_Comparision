from typing import List, Dict, Tuple
from rapidfuzz import fuzz, process
import pandas as pd
from docx import Document
import difflib

FIELDS_KEY = ["topic","subtopic","section_ref","subsection_ref"]

def match_sections(old_units: List[Dict], new_units: List[Dict]) -> List[Dict]:
    # Build maps by composite keys
    def key(u): return (u.get("topic",""), u.get("subtopic",""), u.get("section_ref",""), u.get("subsection_ref",""))
    old_map = { key(u): u for u in old_units }
    new_map = { key(u): u for u in new_units }
    matched = []
    used_new = set()

    # 1) exact key match
    for k, ou in old_map.items():
        if k in new_map:
            nu = new_map[k]
            used_new.add(k)
            status, sim = _status(ou["text"], nu["text"])
            matched.append(_row(ou, nu, status, sim, "exact_key"))
        else:
            matched.append(_row(ou, None, "Removed", 0.0, "unmatched_old"))

    # 2) fuzzy fallback for Removed items by section_ref + heading similarity
    removed_rows = [r for r in matched if r["status"] == "Removed"]
    remaining_new = {k:v for k,v in new_map.items() if k not in used_new}
    for r in removed_rows:
        if not remaining_new:
            break
        # Prefer same section_ref; then compare subtopic text similarity
        candidates = [ (k,v) for k,v in remaining_new.items() if v.get("section_ref")==r.get("old_section_ref") ]
        if not candidates:
            candidates = list(remaining_new.items())
        if not candidates:
            continue
        query = r.get("old_subtopic","") + " " + r.get("old_section_heading","")
        best = None
        best_score = -1
        for k,v in candidates:
            score = fuzz.token_set_ratio(query, (v.get("subtopic","")+" "+v.get("section_heading","")).strip())
            if score > best_score:
                best = (k,v)
                best_score = score
        if best and best_score >= 80:
            kbest, vbest = best
            status, sim = _status(r["old_text"], vbest["text"])
            r.update(_row_updates_from_new(vbest, status, sim, "fuzzy_heading"))
            remaining_new.pop(kbest, None)

    # 3) leftovers are Added
    for k,nu in remaining_new.items():
        matched.append(_row(None, nu, "Added", 0.0, "new_only"))
    return matched

def _status(old: str, new: str) -> Tuple[str, float]:
    if old.strip() == new.strip():
        return "Unchanged", 100.0
    sim = fuzz.token_set_ratio(old, new)
    if sim >= 90: return "Minor edit", float(sim)
    if sim >= 65: return "Modified", float(sim)
    return "Substantially modified", float(sim)

def _row(ou, nu, status, sim, how):
    return {
        "old_topic": ou.get("topic") if ou else None,
        "old_subtopic": ou.get("subtopic") if ou else None,
        "old_section_ref": ou.get("section_ref") if ou else None,
        "old_section_heading": ou.get("section_heading") if ou else None,
        "old_subsection_ref": ou.get("subsection_ref") if ou else None,
        "old_text": ou.get("text") if ou else None,
        "new_topic": nu.get("topic") if nu else None,
        "new_subtopic": nu.get("subtopic") if nu else None,
        "new_section_ref": nu.get("section_ref") if nu else None,
        "new_section_heading": nu.get("section_heading") if nu else None,
        "new_subsection_ref": nu.get("subsection_ref") if nu else None,
        "new_text": nu.get("text") if nu else None,
        "status": status,
        "similarity": sim,
        "match_method": how,
    }

def _row_updates_from_new(nu, status, sim, how):
    return {
        "new_topic": nu.get("topic"),
        "new_subtopic": nu.get("subtopic"),
        "new_section_ref": nu.get("section_ref"),
        "new_section_heading": nu.get("section_heading"),
        "new_subsection_ref": nu.get("subsection_ref"),
        "new_text": nu.get("text"),
        "status": status if status != "Removed" else "Modified",
        "similarity": sim,
        "match_method": how,
    }

def make_excel(matched: List[Dict], path_or_buf):
    df = pd.DataFrame(matched)
    cols = ["status","similarity","match_method",
            "old_topic","old_subtopic","old_section_ref","old_subsection_ref","old_section_heading","old_text",
            "new_topic","new_subtopic","new_section_ref","new_subsection_ref","new_section_heading","new_text"]
    df = df[cols]
    df.to_excel(path_or_buf, index=False)

def make_csv(matched: List[Dict], path_or_buf):
    df = pd.DataFrame(matched)
    cols = ["status","similarity","match_method",
            "old_topic","old_subtopic","old_section_ref","old_subsection_ref","old_section_heading","old_text",
            "new_topic","new_subtopic","new_section_ref","new_subsection_ref","new_section_heading","new_text"]
    df = df[cols]
    df.to_csv(path_or_buf, index=False)

def make_word(matched: List[Dict], path_or_buf, title="Act Comparison Report"):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph("This report summarizes the differences between the Old Act and the New Act, organized by Topic/SubTopic/Section/Subsection.")

    total = len(matched)
    added = sum(1 for r in matched if r["status"] == "Added")
    removed = sum(1 for r in matched if r["status"] == "Removed")
    modified = sum(1 for r in matched if "Modified" in r["status"] or "Minor" in r["status"] or r["status"] == "Substantially modified")
    unchanged = sum(1 for r in matched if r["status"] == "Unchanged")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(f"Total compared units: {total}")
    doc.add_paragraph(f"Added: {added}, Removed: {removed}, Modified: {modified}, Unchanged: {unchanged}")

    doc.add_heading("Detailed Narrative", level=1)
    for r in matched:
        heading = (r.get("new_section_heading") or r.get("old_section_heading") or "").strip() or "(Untitled)"
        path = " > ".join([p for p in [r.get("new_topic") or r.get("old_topic"),
                                       r.get("new_subtopic") or r.get("old_subtopic"),
                                       r.get("new_section_ref") or r.get("old_section_ref"),
                                       r.get("new_subsection_ref") or r.get("old_subsection_ref")] if p])
        p = doc.add_paragraph()
        p.add_run(f"{path} — {r['status']} (Similarity: {int(r['similarity'])}%)").bold = True
        if r["status"] == "Added":
            doc.add_paragraph(r.get("new_text") or "")
        elif r["status"] == "Removed":
            doc.add_paragraph(r.get("old_text") or "")
        else:
            doc.add_paragraph("Old:"); doc.add_paragraph(r.get("old_text") or "")
            doc.add_paragraph("New:"); doc.add_paragraph(r.get("new_text") or "")
        doc.add_paragraph("")
    doc.save(path_or_buf)

def inline_diff(old: str, new: str) -> str:
    sm = difflib.SequenceMatcher(None, old.split(), new.split())
    out = []
    old_words = old.split()
    new_words = new.split()
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            out.extend(old_words[i1:i2])
        elif tag == 'replace':
            out.append('<del>' + ' '.join(old_words[i1:i2]) + '</del>')
            out.append('<ins>' + ' '.join(new_words[j1:j2]) + '</ins>')
        elif tag == 'delete':
            out.append('<del>' + ' '.join(old_words[i1:i2]) + '</del>')
        elif tag == 'insert':
            out.append('<ins>' + ' '.join(new_words[j1:j2]) + '</ins>')
    return ' '.join(out)
